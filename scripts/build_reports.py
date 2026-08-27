"""Build the monthly deliverables from correlate.py output.

  entra-pim-correlation-YYYY-MM.xlsx      findings workbook
  entra-pim-review-summary-YYYY-MM.docx   one-page executive summary

Workbook conventions (CLAUDE.md): Arial throughout, header row frozen, autofilter on,
no merged cells in data ranges. Summary figures are live formulas over the data sheets,
never Python-computed literals, so the workbook recalculates if a sheet is edited.
Any figure that cannot be a formula (raw source row counts) is labelled with the source
file it came from, and the Evidence sheet carries the manifest for every input.

Uses python-docx rather than docx-js so the whole workflow is a single Python toolchain
on the reviewer's machine.
"""

from __future__ import annotations

import argparse
import json
import sys
from datetime import datetime, timezone
from pathlib import Path

import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

sys.path.insert(0, str(Path(__file__).resolve().parent))

from common import month_paths

FONT = "Arial"
HDR_FILL = PatternFill("solid", fgColor="1F3864")
HDR_FONT = Font(name=FONT, bold=True, color="FFFFFF", size=10)
BODY_FONT = Font(name=FONT, size=10)
TITLE_FONT = Font(name=FONT, bold=True, size=14)
LABEL_FONT = Font(name=FONT, bold=True, size=10)
NOTE_FONT = Font(name=FONT, italic=True, size=9, color="595959")
SEV_FILL = {"High": PatternFill("solid", fgColor="F8CBAD"),
            "Medium": PatternFill("solid", fgColor="FFE699"),
            "Low": PatternFill("solid", fgColor="E2EFDA"),
            "Informational": PatternFill("solid", fgColor="DDEBF7")}
THIN = Side(style="thin", color="BFBFBF")

ROW_BUFFER = 100  # formula ranges extend this far past the data so edits still recalculate

# --------------------------------------------------------------- advisory triage columns
#
# These three Decisions columns group the exception list so a role owner is not handed
# hundreds of undifferentiated rows. They are ADVISORY and carry no authority: they set no
# severity, change no count, and leave decision / decided_by / decision_date untouched.
# Per CLAUDE.md the keep / modify / revoke decision belongs to the role owner.

REVIEW_THEME = {
    "permanent_assignment_outside_pim": "Standing access outside time-binding",
    "permanent_eligibility_grant": "Standing access outside time-binding",
    "uncovered_privileged_action": "Action outside any activation window",
    "global_administrator_activation": "Highest-privilege role usage",
    "failed_timebound_assignment_request": "Failing control - configuration",
    "failed_activation": "Failing control - configuration",
    "weak_justification": "Justification quality",
    "missing_justification": "Justification quality",
    "activation_no_actions": "Activation with no observed use",
    "activation_request_not_completed": "Activation with no observed use",
    "activation_on_behalf_of_other": "Delegation path",
    "eligibility_grant_for_review": "Eligibility granted",
    "off_hours_activation": "Timing - check reporting_timezone first",
    "break_glass_activity": "Break-glass - reported, not for revocation",
}

ADVISORY = {
    "permanent_assignment_outside_pim":
        "Confirm this was approved. If it stands, convert to a PIM-managed, time-bound "
        "assignment; a permanent grant made outside PIM is subject to no activation, "
        "approval or expiry.",
    "permanent_eligibility_grant":
        "Confirm approved, then convert to time-bound eligibility. Permanent eligibility "
        "defeats the purpose of PIM.",
    "uncovered_privileged_action":
        "Establish how the actor held this privilege with no activation covering it - "
        "standing assignment, PIM bypass, or attribution falling outside the fixed window.",
    "global_administrator_activation":
        "Confirm Global Administrator was necessary and no narrower role would serve.",
    "failed_timebound_assignment_request":
        "Diagnose the block rather than attest: every such request failed this period, "
        "which indicates configuration or policy, not user error.",
    "failed_activation":
        "Diagnose the cause; repeated failures may indicate misconfiguration.",
    "weak_justification":
        "Assess before actioning - a bare ticket reference may be adequate evidence while "
        "still failing the character-count rule. Consider tuning the rule.",
    "missing_justification":
        "Require a justification for future activations of this role.",
    "activation_no_actions":
        "Candidate for removal, but read alongside the fixed attribution window: an admin "
        "who worked outside that window looks identical to one who did nothing.",
    "activation_on_behalf_of_other":
        "Confirm this is an approved delegation path.",
    "eligibility_grant_for_review":
        "Confirm the grant was approved and is time-bound.",
    "off_hours_activation":
        "Confirm reporting_timezone matches the operating timezone before treating this as "
        "a finding; a UTC default inflates off-hours counts for a non-UTC organisation.",
    "break_glass_activity":
        "Reported for visibility only. Not for revocation.",
}

READ_ONLY_RIDER = (" NOTE: this is a read/telemetry operation, which is weak evidence of "
                   "standing privileged access - triage below the write operations in this "
                   "class.")

# Audit activities that read state rather than change it. Anything not listed is treated as
# a write or unclear, so an unrecognised activity is never wrongly discounted.
READ_ONLY_ACTIVITIES = {
    "groupsodatav4_get", "validate user authentication",
    "group_getdynamicgroupproperties", "settings_getsettingsasync",
    "approval_getall", "get authenticationeventlisteners",
}


def classify_action(exception_class: str, detail: str) -> str:
    """read-only / write-or-unclear for an uncovered action; blank for other classes.

    Informational only - it does not alter severity, so approved counts stay valid.
    """
    if exception_class != "uncovered_privileged_action":
        return ""
    start = str(detail).find("'")
    end = str(detail).find("'", start + 1)
    if start < 0 or end < 0:
        return "write or unclear"
    activity = str(detail)[start + 1:end].strip().casefold()
    return "read-only" if activity in READ_ONLY_ACTIVITIES else "write or unclear"


def last_row(df: pd.DataFrame) -> int:
    """Formula range end: data extent plus a small buffer. Deliberately not a 100k-row
    ceiling - open-ended ranges across eight sheets make recalculation crawl."""
    return max(2, len(df) + 1) + ROW_BUFFER


def col_of(df: pd.DataFrame, header: str) -> str:
    """Excel column letter for a dataframe header (1-based, data starts row 2)."""
    return get_column_letter(list(df.columns).index(header) + 1)


def write_sheet(wb: Workbook, name: str, df: pd.DataFrame, note: str = "") -> None:
    ws = wb.create_sheet(name)
    if df is None or df.empty:
        ws["A1"] = f"No rows for this period."
        ws["A1"].font = BODY_FONT
        if note:
            ws["A2"] = note
            ws["A2"].font = NOTE_FONT
        return

    ws.append(list(df.columns))
    for cell in ws[1]:
        cell.font, cell.fill = HDR_FONT, HDR_FILL
        cell.alignment = Alignment(vertical="center", wrap_text=True)
    ws.row_dimensions[1].height = 28

    sev_col = col_of(df, "severity") if "severity" in df.columns else None
    for rec in df.itertuples(index=False):
        ws.append(["" if pd.isna(v) else v for v in rec])
        r = ws.max_row
        for cell in ws[r]:
            cell.font = BODY_FONT
            cell.alignment = Alignment(vertical="top", wrap_text=False)
            cell.border = Border(bottom=THIN)
        if sev_col:
            fill = SEV_FILL.get(str(ws[f"{sev_col}{r}"].value))
            if fill:
                ws[f"{sev_col}{r}"].fill = fill

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:{get_column_letter(len(df.columns))}{ws.max_row}"
    for i, colname in enumerate(df.columns, start=1):
        width = max(len(str(colname)) + 3,
                    min(60, int(df[colname].astype(str).str.len().quantile(0.95)) + 3))
        ws.column_dimensions[get_column_letter(i)].width = width


def build_summary(ws, stats: dict, acts: pd.DataFrame, exc: pd.DataFrame,
                  corr: pd.DataFrame, unc: pd.DataFrame) -> None:
    month = stats.get("month", "")
    audit_ok = bool(stats.get("audit_available"))
    pim_file = stats.get("pim_file") or "(none)"
    audit_file = stats.get("audit_file") or "(not provided)"

    ws["A1"] = f"Entra PIM / PAM Monthly Review - {month}"
    ws["A1"].font = TITLE_FONT
    ws["A2"] = (f"Generated {datetime.now(timezone.utc):%Y-%m-%d %H:%M} UTC   |   "
                f"correlation window {stats.get('activation_window_hours')}h   |   "
                f"reporting timezone {stats.get('reporting_timezone')}")
    ws["A2"].font = NOTE_FONT

    row = 4

    def section(title):
        nonlocal row
        ws[f"A{row}"] = title
        ws[f"A{row}"].font = Font(name=FONT, bold=True, size=11, color="1F3864")
        row += 1

    def line(label, value, note=""):
        nonlocal row
        ws[f"A{row}"] = label
        ws[f"A{row}"].font = LABEL_FONT
        ws[f"B{row}"] = value
        ws[f"B{row}"].font = BODY_FONT
        if note:
            ws[f"C{row}"] = note
            ws[f"C{row}"].font = NOTE_FONT
        row += 1

    A_END, E_END, C_END, U_END = (last_row(acts), last_row(exc), last_row(corr), last_row(unc))
    a_id = col_of(acts, "activation_id") if not acts.empty else "A"
    a_status = col_of(acts, "correlation_status") if not acts.empty else "A"
    a_role = col_of(acts, "Entra Role") if not acts.empty else "A"
    e_cls = col_of(exc, "exception_class") if not exc.empty else "B"
    e_sev = col_of(exc, "severity") if not exc.empty else "C"

    section("Scope and source data")
    line("PIM activity file", pim_file, "raw source, unmodified in input/")
    line("Directory audit file", audit_file,
         "" if audit_ok else "MISSING - see limitation note below")
    line("PIM rows as exported", stats.get("pim_rows_raw", 0),
         f"from {pim_file}; see Evidence sheet")
    line("Exact duplicate rows dropped", stats.get("pim_exact_duplicates_dropped", 0),
         "log-platform fan-out; every copy identical")
    line("Distinct PIM events analysed", stats.get("pim_rows_deduped", 0), "")
    line("Period observed in data",
         f"{str(stats.get('pim_period_observed_start'))[:10]} to "
         f"{str(stats.get('pim_period_observed_end'))[:10]}", "")
    line("Distinct privileged actors", stats.get("pim_distinct_actors", 0), "")
    row += 1

    section("Activations")
    line("Successful activations", f"=COUNTA(Activations!${a_id}$2:${a_id}${A_END})",
         "one row per activation on the Activations sheet")
    line("Activation requests recorded", stats.get("activation_requests", 0), "")
    line("Self-activations", stats.get("activations_self", 0), "normal PIM behaviour")
    line("Activated for another account", stats.get("activations_on_behalf_of_other", 0),
         "delegation path - confirm approved")
    line("Global Administrator activations",
         f'=COUNTIF(Activations!${a_role}$2:${a_role}${A_END},"Global Administrator")',
         "highest-privilege role")
    row += 1

    section("Correlation outcome")
    if audit_ok:
        line("Audit rows eligible for correlation",
             stats.get("audit_rows_eligible_for_correlation", 0),
             "PIM's own events excluded so an activation cannot match itself")
        line("Correlated action rows", f"=COUNTA('Correlated Actions'!$A$2:$A${C_END})",
             "activation-to-action pairs")
        line("Distinct audit events matched", stats.get("correlated_distinct_events", 0), "")
        line("Ambiguous attributions", stats.get("correlated_ambiguous_rows", 0),
             "event fell inside two overlapping windows for the same actor")
        line("Activations with observed actions",
             f'=COUNTIF(Activations!${a_status}$2:${a_status}${A_END},"actions observed")', "")
        line("Activations with no observed actions",
             f'=COUNTIF(Activations!${a_status}$2:${a_status}${A_END},"no actions observed")',
             "candidates for removal")
        line("Actions outside any activation window",
             f"=COUNTA('Uncovered Actions'!$A$2:$A${U_END})",
             "possible standing access or PIM bypass")
    else:
        line("Correlation status", "NOT PERFORMED",
             "no directory audit file for this period")
        line("Activations classified", "unknown - no audit data",
             "absence of audit data is not evidence that no action was taken")
    row += 1

    section("Exceptions raised")
    line("Total exceptions", f"=COUNTA(Exceptions!$A$2:$A${E_END})", "")
    for sev in ("High", "Medium", "Low", "Informational"):
        line(f"  {sev}", f'=COUNTIF(Exceptions!${e_sev}$2:${e_sev}${E_END},"{sev}")', "")
    row += 1
    if not exc.empty:
        ws[f"A{row}"] = "By class"
        ws[f"A{row}"].font = LABEL_FONT
        row += 1
        for cls in sorted(exc["exception_class"].unique()):
            ws[f"A{row}"] = f"  {cls}"
            ws[f"A{row}"].font = BODY_FONT
            ws[f"B{row}"] = f'=COUNTIF(Exceptions!${e_cls}$2:${e_cls}${E_END},"{cls}")'
            ws[f"B{row}"].font = BODY_FONT
            row += 1
    row += 1

    section("Limitations to read alongside these figures")
    notes = [
        f"Attribution uses a fixed {stats.get('activation_window_hours')}-hour window from each "
        f"activation. Real start/end times are available from Graph "
        f"roleManagement/directory/roleAssignmentScheduleRequests and would remove this "
        f"approximation.",
        "Correlation is temporal, not causal: an action inside a window is not proof the "
        "activated role authorised it.",
        "Recommendations are advisory. Keep / modify / revoke decisions belong to the role owner.",
    ]
    if not audit_ok:
        notes.insert(0, "No directory audit export was available, so no activation can be "
                        "reported as unused. Those rows read 'unknown - no audit data'.")
    if not stats.get("break_glass_declared"):
        notes.append("No break-glass accounts are declared in config, so any such account will "
                     "appear as an uncovered-action finding.")
    if stats.get("audit_chunk_rows_collapsed"):
        notes.append(
            f"{stats['audit_chunk_fragment_rows']} audit rows were fragments of "
            f"{stats['audit_chunked_events']} events split by Graph across rows, and were "
            f"reassembled; counting them separately would have inflated action volume by "
            f"{stats['audit_chunk_rows_collapsed']} rows.")
    if stats.get("audit_content_identical_rows"):
        notes.append(
            f"{stats['audit_content_identical_rows']} audit rows are indistinguishable from "
            f"another row in every exported field except the sub-second timestamp. They are "
            f"kept rather than deduplicated, so action volume is an upper bound for those "
            f"events.")
    if stats.get("pim_unmapped_action_rows"):
        notes.append(
            f"{stats['pim_unmapped_action_rows']} PIM rows carry an event type this workflow "
            f"does not map and therefore generate no findings: "
            f"{', '.join(f'{k} ({v})' for k, v in sorted((stats.get('pim_unmapped_actions') or {}).items()))}.")
    for n in notes:
        ws[f"A{row}"] = f"- {n}"
        ws[f"A{row}"].font = NOTE_FONT
        ws[f"A{row}"].alignment = Alignment(wrap_text=True, vertical="top")
        ws.row_dimensions[row].height = 26
        row += 1

    ws.column_dimensions["A"].width = 42
    ws.column_dimensions["B"].width = 26
    ws.column_dimensions["C"].width = 62
    ws.sheet_view.showGridLines = False


def build_evidence(ws, stats: dict) -> None:
    ws["A1"] = "Evidence - provenance for every figure in this workbook"
    ws["A1"].font = TITLE_FONT
    headers = ["file", "source", "endpoint", "filter", "period", "row_count",
               "sha256", "exported_utc", "notes"]
    ws.append([])
    ws.append(headers)
    for cell in ws[3]:
        cell.font, cell.fill = HDR_FONT, HDR_FILL
    entries = (stats.get("manifest") or {}).get("entries") or []
    if not entries:
        ws.append(["(no export-manifest.json present - files were placed manually)"])
        ws[4][0].font = NOTE_FONT
    for e in entries:
        ws.append([e.get(h, "") for h in headers])
        for cell in ws[ws.max_row]:
            cell.font = BODY_FONT
    ws.freeze_panes = "A4"
    for i, w in enumerate([40, 20, 40, 52, 12, 12, 30, 26, 60], start=1):
        ws.column_dimensions[get_column_letter(i)].width = w


def build_docx(path: Path, stats: dict, acts: pd.DataFrame, exc: pd.DataFrame,
               unc: pd.DataFrame) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    audit_ok = bool(stats.get("audit_available"))
    month = stats.get("month", "")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name, style.font.size = FONT, Pt(10)

    h = doc.add_heading(f"Entra PIM / PAM Monthly Review - {month}", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph()
    r = sub.add_run(f"Prepared {datetime.now(timezone.utc):%Y-%m-%d} UTC  |  "
                    f"Owner: Vineet Gupta (vgupta@caqh.org), CAQH  |  "
                    f"Correlation window: {stats.get('activation_window_hours')}h")
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_heading("Scope", level=1)
    observed = (f"{str(stats.get('pim_period_observed_start'))[:10]} to "
                f"{str(stats.get('pim_period_observed_end'))[:10]}")
    doc.add_paragraph(
        f"PIM activity and Entra directory audit events for {month}, correlated to establish "
        f"what each administrator did with privilege after activating it. "
        f"Events observed in the source data span {observed}. "
        f"Source files: {stats.get('pim_file')}; {stats.get('audit_file') or 'no audit export'}. "
        f"{stats.get('pim_rows_raw', 0)} exported PIM rows reduced to "
        f"{stats.get('pim_rows_deduped', 0)} distinct events after removing "
        f"{stats.get('pim_exact_duplicates_dropped', 0)} identical duplicate copies.")

    doc.add_heading("Counts", level=1)
    rows = [
        ("Distinct PIM events analysed", stats.get("pim_rows_deduped", 0)),
        ("Successful activations", stats.get("activations_successful", 0)),
        ("Distinct privileged actors", stats.get("pim_distinct_actors", 0)),
        ("Global Administrator activations",
         int((acts["Entra Role"] == "Global Administrator").sum()) if not acts.empty else 0),
        ("Exceptions raised", len(exc)),
        ("  High severity", int((exc["severity"] == "High").sum()) if not exc.empty else 0),
        ("  Medium severity", int((exc["severity"] == "Medium").sum()) if not exc.empty else 0),
        ("Actions correlated to an activation",
         stats.get("correlated_action_rows", 0) if audit_ok else "n/a - no audit data"),
        ("Actions outside any activation window",
         stats.get("uncovered_action_rows", 0) if audit_ok else "n/a - no audit data"),
        ("Revoked this cycle", "0 - decisions pending role-owner attestation"),
        ("Retained this cycle", "0 - decisions pending role-owner attestation"),
    ]
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    for label, val in rows:
        c = t.add_row().cells
        c[0].text, c[1].text = str(label), str(val)
        for cell in c:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size, run.font.name = Pt(9), FONT

    doc.add_heading("Notable risks", level=1)
    risks = []
    if not exc.empty:
        for cls, n in exc[exc["severity"] == "High"]["exception_class"].value_counts().items():
            risks.append(f"{n} x {cls.replace('_', ' ')} - highest priority for this cycle.")
    ga = int((acts["Entra Role"] == "Global Administrator").sum()) if not acts.empty else 0
    if ga and stats.get("activations_successful"):
        pct = round(100 * ga / stats["activations_successful"])
        risks.append(f"Global Administrator accounts for {pct}% of activations ({ga} of "
                     f"{stats['activations_successful']}). Consider whether a lower-privilege "
                     f"role would serve these tasks.")
    if stats.get("timebound_requests_failed"):
        risks.append(f"{stats['timebound_requests_failed']} of "
                     f"{stats.get('timebound_requests_total')} time-bound assignment requests "
                     f"failed. A complete or near-complete failure rate points at policy or "
                     f"configuration, not user error.")
    if not audit_ok:
        risks.append("No directory audit export was available, so unused activations could not "
                     "be identified this cycle. This is the review's main gap.")
    if not risks:
        risks.append("No high-severity exceptions were raised for this period.")
    for rk in risks:
        doc.add_paragraph(rk, style="List Bullet")

    doc.add_heading("Remediation status", level=1)
    doc.add_paragraph(
        "All exceptions are open pending role-owner attestation. Each is routed for a keep, "
        "modify, or revoke decision; the decision, decider, and date are recorded on the "
        "Decisions tab of the findings workbook. Unresolved items carry into next month's cycle.")

    doc.add_heading("Limitations", level=1)
    lim = [f"Attribution uses a fixed {stats.get('activation_window_hours')}-hour window from "
           f"each activation rather than the real activation expiry, which over-attributes when "
           f"an admin deactivates early and under-attributes if the configured maximum is longer.",
           "Correlation is temporal, not causal. An action inside a window is not proof that the "
           "activated role authorised it.",
           "Every figure traces to a source file listed on the workbook's Evidence sheet. "
           "Recommendations are advisory; revocation decisions belong to the role owner."]
    if stats.get("audit_content_identical_rows"):
        lim.append(
            f"{stats['audit_content_identical_rows']} audit rows cannot be distinguished from "
            f"another row by any exported field except a sub-second timestamp difference. They "
            f"are retained rather than deduplicated, because the field that would separate them "
            f"is absent from the export; action volume for those events is an upper bound.")
    if not audit_ok:
        lim.insert(0, "No directory audit data was available for this period, so no activation "
                      "is reported as unused. Absence of evidence is not evidence of absence.")
    for l in lim:
        doc.add_paragraph(l, style="List Bullet")

    doc.save(path)


def main() -> int:
    ap = argparse.ArgumentParser(description="Build xlsx + docx deliverables.")
    ap.add_argument("--month", required=True)
    ap.add_argument("--run", type=int, default=1)
    ap.add_argument("--label", default=None)
    args = ap.parse_args()

    paths = month_paths(args.month, args.run)
    label = args.label or args.month
    out = paths["output"]

    stats_path = out / f"correlation-stats-{label}.json"
    if not stats_path.exists():
        print(f"Missing {stats_path.name} - run correlate.py first.", file=sys.stderr)
        return 2
    stats = json.loads(stats_path.read_text(encoding="utf-8"))

    def read(name):
        p = out / f"{name}-{label}.csv"
        return pd.read_csv(p, keep_default_na=False) if p.exists() and p.stat().st_size > 2 \
            else pd.DataFrame()

    acts, corr, unc, exc = (read("activations"), read("correlated-actions"),
                            read("uncovered-actions"), read("exceptions"))

    wb = Workbook()
    wb.remove(wb.active)
    build_summary(wb.create_sheet("Summary"), stats, acts, exc, corr, unc)
    write_sheet(wb, "Activations", acts)
    write_sheet(wb, "Correlated Actions", corr,
                "No audit export for this period." if not stats.get("audit_available") else "")
    unmatched = acts[acts["correlation_status"] == "no actions observed"] \
        if not acts.empty and "correlation_status" in acts else pd.DataFrame()
    write_sheet(wb, "Unmatched Activations", unmatched,
                "Requires a directory audit export to populate."
                if not stats.get("audit_available") else "")
    write_sheet(wb, "Uncovered Actions", unc,
                "Requires a directory audit export to populate."
                if not stats.get("audit_available") else "")
    write_sheet(wb, "Exceptions", exc)
    decisions = pd.DataFrame(columns=["exception_id", "exception_class", "actor", "entra_role",
                                      "review_theme", "action_type",
                                      "advisory_recommendation",
                                      "decision (keep/modify/revoke)", "decided_by",
                                      "decision_date", "target_remediation_date", "notes"])
    if not exc.empty:
        decisions = exc[["exception_id", "exception_class", "actor", "entra_role"]].copy()
        decisions["review_theme"] = exc["exception_class"].map(
            lambda c: REVIEW_THEME.get(c, "Unthemed - triage manually"))
        decisions["action_type"] = [classify_action(cls, det) for cls, det
                                    in zip(exc["exception_class"], exc["detail"])]
        decisions["advisory_recommendation"] = [
            ADVISORY.get(cls, "Route to role owner for keep / modify / revoke.")
            + (READ_ONLY_RIDER if at == "read-only" else "")
            for cls, at in zip(exc["exception_class"], decisions["action_type"])]
        for c in ("decision (keep/modify/revoke)", "decided_by", "decision_date",
                  "target_remediation_date", "notes"):
            decisions[c] = ""
    write_sheet(wb, "Decisions", decisions,
                "review_theme, action_type and advisory_recommendation are advisory only "
                "and set no severity. Role owner fills the blank columns: decision, "
                "decider, date.")
    build_evidence(wb.create_sheet("Evidence"), stats)

    xlsx = out / f"entra-pim-correlation-{label}.xlsx"
    wb.save(xlsx)

    docx = out / f"entra-pim-review-summary-{label}.docx"
    build_docx(docx, stats, acts, exc, unc)

    print(f"  wrote {xlsx.name} ({len(wb.sheetnames)} sheets)")
    print(f"  wrote {docx.name}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
